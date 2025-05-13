import os
import shutil
import random
from docx import Document
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox


ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")


def center_window(win, width=None, height=None):
    win.update_idletasks()
    screen_width = win.winfo_screenwidth()
    screen_height = win.winfo_screenheight()

    if width is None or height is None:
        current_geometry = win.geometry()
        wh_part = current_geometry.split('+')[0]
        w_str, h_str = wh_part.split('x')
        width = int(w_str)
        height = int(h_str)

    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    win.geometry(f"{width}x{height}+{x}+{y}")


root = ctk.CTk()
root.title("Тестирование")
root.geometry("900x600")
center_window(root, 600, 400)


def close_app():
    if messagebox.askyesno("Подтверждение", "Вы действительно хотите выйти?"):
        # Удаляем временные изображения
        temp_img_dir = "temp_images"
        if os.path.exists(temp_img_dir):
            shutil.rmtree(temp_img_dir)
        root.destroy()


root.protocol("WM_DELETE_WINDOW", close_app)


def start_test_with_incorrect_questions():
    file_path = filedialog.askopenfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        title="Выберите файл с вопросами для теста"
    )
    if not file_path:
        return  # Пользователь закрыл диалог

    try:
        questions = parse_questions(file_path)
        if not questions:
            raise ValueError("Файл не содержит вопросов или он неверно отформатирован.")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось загрузить вопросы: {e}")
        return

    # Проверка перемешивания
    shuffle_questions = shuffle_answers_var.get()

    # Если нужно перемешивать вопросы
    if shuffle_questions:
        questions = random.sample(questions, len(questions))

    # Запуск теста с загруженными вопросами
    root.withdraw()
    TestWindow(questions, shuffle_answers_var.get(), 50)


def parse_questions(file_name):
    document = Document(file_name)
    questions = []
    temp_img_dir = "temp_images"
    if not os.path.exists(temp_img_dir):
        os.makedirs(temp_img_dir)
    img_counter = 0
    paragraphs = document.paragraphs
    
    i = 0
    while i < len(paragraphs):
        current_para_obj = paragraphs[i]
        current_para_full_text = current_para_obj.text 

        # Используем lstrip(), чтобы тег считывался, даже если перед ним есть пробелы в начале абзаца
        if current_para_full_text.lstrip().startswith("<question>"):
            
            # Инициализация для текущего вопроса
            parsed_q_text = ""
            current_question_variants = []
            image_path = None # Сбрасываем для каждого нового вопроса

            # --- Обработка текущего абзаца (который содержит <question>) ---
            # Этот абзац может содержать текст вопроса и варианты ответов.
            lines_in_current_para = current_para_full_text.splitlines()
            
            is_parsing_question_text = False # Флаг, что мы в блоке текста вопроса
            question_text_lines_buffer = []

            for line_content in lines_in_current_para:
                stripped_line = line_content.strip()
                
                if stripped_line.startswith("<question>"):
                    # Начало блока вопроса
                    text_after_tag = stripped_line.replace("<question>", "").strip()
                    if text_after_tag: # Если текст есть на той же строке, что и тег
                        question_text_lines_buffer.append(text_after_tag)
                    is_parsing_question_text = True # Активируем сбор текста вопроса
                    # Продолжаем, чтобы собрать многострочный текст вопроса, если он есть
                
                elif is_parsing_question_text:
                    # Мы находимся в процессе сбора текста вопроса
                    if stripped_line.startswith("<variant>"):
                        # Встретили <variant> внутри блока вопроса - текст вопроса закончился
                        is_parsing_question_text = False
                        # Этот вариант принадлежит текущему вопросу
                        current_question_variants.append(stripped_line.replace("<variant>", "").strip())
                    elif stripped_line: # Если строка не пустая и не <variant>
                        # Это продолжение многострочного текста вопроса
                        question_text_lines_buffer.append(stripped_line)
                
                elif stripped_line.startswith("<variant>"):
                    # Мы уже вышли из блока текста вопроса (или он был однострочным)
                    # и это строка с вариантом в том же абзаце
                    current_question_variants.append(stripped_line.replace("<variant>", "").strip())
            
            parsed_q_text = "\n".join(question_text_lines_buffer).strip()

            if not parsed_q_text: # Если тег <question> был, но текста вопроса не нашлось
                i += 1 # Пропускаем этот "пустой" вопрос
                continue

            # --- Поиск картинки и доп. вариантов в СЛЕДУЮЩИХ абзацах ---
            # Индекс для начала поиска в следующих абзацах
            next_para_scan_idx = i + 1
            
            # Индекс, с которого начнется сбор доп. вариантов (после картинки или сразу)
            start_further_variants_idx = next_para_scan_idx
            
            # Фаза 1: Поиск картинки в следующих абзацах
            img_search_runner_idx = next_para_scan_idx
            while img_search_runner_idx < len(paragraphs):
                img_candidate_para_obj = paragraphs[img_search_runner_idx]
                img_candidate_para_text = img_candidate_para_obj.text.strip()

                if img_candidate_para_text.startswith("<variant>") or \
                   img_candidate_para_text.startswith("<question>"):
                    # Дошли до вариантов или нового вопроса, картинки для текущего вопроса здесь нет
                    start_further_variants_idx = img_search_runner_idx
                    break 
                
                found_image_in_para = False
                for run in img_candidate_para_obj.runs:
                    if 'graphic' in run._element.xml:
                        drawing = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        if drawing:
                            embed_id = drawing[0].attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                rel = document.part.rels.get(embed_id)
                                if rel and "image" in rel.reltype:
                                    img_data = rel.target_part.blob
                                    img_counter += 1
                                    image_path = os.path.join(temp_img_dir, f"question_img_{img_counter}.png")
                                    with open(image_path, "wb") as img_file:
                                        img_file.write(img_data)
                                    found_image_in_para = True
                                    break # Выход из цикла по runs
                if found_image_in_para:
                    start_further_variants_idx = img_search_runner_idx + 1 # Варианты ищем ПОСЛЕ картинки
                    break # Выход из цикла поиска картинки
                
                img_search_runner_idx += 1
            else: # Цикл поиска картинки завершился без break (дошли до конца документа)
                start_further_variants_idx = img_search_runner_idx

            # Фаза 2: Сбор ДОПОЛНИТЕЛЬНЫХ вариантов из следующих абзацев
            # `next_question_or_eof_idx` будет указывать на начало следующего вопроса или конец файла
            next_question_or_eof_idx = start_further_variants_idx
            
            further_variants_runner_idx = start_further_variants_idx
            while further_variants_runner_idx < len(paragraphs):
                further_variant_para_obj = paragraphs[further_variants_runner_idx]
                further_variant_para_text = further_variant_para_obj.text.strip()

                if further_variant_para_text.startswith("<variant>"):
                    current_question_variants.append(further_variant_para_text.replace("<variant>", "").strip())
                elif further_variant_para_text.startswith("<question>"):
                    next_question_or_eof_idx = further_variants_runner_idx # Следующий вопрос найден
                    break
                # Игнорируем другие строки между вариантами в последующих абзацах
                
                further_variants_runner_idx += 1
            else: # Цикл сбора доп. вариантов завершился без break (конец документа)
                next_question_or_eof_idx = further_variants_runner_idx
            
            questions.append({
                "question": parsed_q_text,
                "variants": current_question_variants,
                "correct_index": 0, # По-прежнему предполагаем, что первый вариант правильный
                "image": image_path
            })
            
            # Устанавливаем `i` так, чтобы следующая итерация основного цикла
            # началась с абзаца, где был найден следующий тег <question> (или с конца документа)
            i = next_question_or_eof_idx - 1 
        
        i += 1 # Переход к следующему абзацу в основном цикле
        
    return questions


def start_test():
    try:
        time_limit_str = time_limit_var.get().strip()
        if not time_limit_str.isdigit():
            time_limit = 50
        else:
            time_limit = int(time_limit_str)
        num_questions_str = question_count_var.get().strip()
        if not num_questions_str.isdigit():
            raise ValueError("Введите число.")
        num_questions = int(num_questions_str)
        if num_questions <= 0:
            raise ValueError("Количество вопросов должно быть положительным числом.")

        questions = parse_questions("URBU.docx")
        if not questions:
            raise ValueError("В файле нет вопросов или они неверно отформатированы.")
    except ValueError as ve:
        messagebox.showerror("Ошибка", f"Некорректное значение: {ve}")
        return
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось загрузить вопросы: {e}")
        return

    shuffled_questions = random.sample(questions, k=min(num_questions, len(questions)))
    root.withdraw()
    TestWindow(shuffled_questions, shuffle_answers_var.get(), time_limit)


def export_incorrect_questions_as_original_format(results):
    incorrect_results = [r for r in results if not r["is_correct"]]
    if not incorrect_results:
        messagebox.showinfo("Информация", "Нет неправильных ответов для экспорта.")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        title="Сохранить файл как"
    )
    if not file_path:
        return

    document = Document()

    for result in incorrect_results:
        document.add_paragraph(f"<question> {result['question']}")

       
        image_path = result.get('image') 
        if image_path and os.path.exists(image_path):
            try:
                # from docx.shared import Inches
                # document.add_picture(image_path, width=Inches(4.0))
                document.add_picture(image_path)
                # Можно добавить пустой параграф после картинки для разделения
                # document.add_paragraph()
            except Exception as e:
                print(f"Не удалось добавить картинку {image_path} в документ: {e}")
                # Можно добавить плейсхолдер в документ
                document.add_paragraph(f"[Не удалось вставить изображение: {os.path.basename(image_path)}]")

        for variant in result["variants"]:
            document.add_paragraph(f"<variant> {variant}")
        # Добавляем пустой абзац для разделения вопросов, если хотите
        # document.add_paragraph()

    try:
        document.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Файл успешно сохранен: {file_path}")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {e}")

def export_incorrect_answers(results):
    # Фильтруем только неправильные ответы
    incorrect_results = [r for r in results if not r["is_correct"]]
    if not incorrect_results:
        messagebox.showinfo("Информация", "Нет неправильных ответов для экспорта.")
        return

    # Открываем файловый диалог для выбора пути сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        title="Сохранить файл как"
    )
    if not file_path:  # Если пользователь закрыл диалог без выбора
        return

    # Создаем документ Word
    document = Document()
    document.add_heading("Неправильные ответы", level=1)

    for result in incorrect_results:
        document.add_heading(result["question"], level=2)
        document.add_paragraph(f"{result['correct']}")

    # Сохраняем файл по указанному пути
    try:
        document.save(file_path)
        messagebox.showinfo("Экспорт завершен", f"Файл успешно сохранен: {file_path}")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {e}")


class TestWindow(ctk.CTkToplevel):
    def __init__(self, questions, shuffle_answers, time_limit):
        super().__init__()
        self.questions = questions
        self.shuffle_answers = shuffle_answers
        self.current_question = 0
        self.score = 0
        self.results = []
        self.time_remaining = time_limit * 60

        self.title("Тестирование")
        self.state('zoomed')

        self.protocol("WM_DELETE_WINDOW", self.close_test)
        self.create_widgets()
        self.start_timer()
        self.show_question()

    def create_widgets(self):
        self.main_frame = ctk.CTkFrame(self, corner_radius=15)
        self.main_frame.pack(expand=True, fill="both", padx=20, pady=20)

        self.timer_label = ctk.CTkLabel(self.main_frame, text="", font=("Arial", 20, "bold"), text_color="red")
        self.timer_label.pack(pady=5, fill="x", expand=True)

        self.question_label = ctk.CTkLabel(self.main_frame, text="", wraplength=1000, justify="left",
                                           font=("Arial", 14, "bold"))
        self.question_label.pack(pady=10, fill="x", expand=True)

        self.image_label = ctk.CTkLabel(self.main_frame, text="")
        self.image_label.pack(pady=5)

        self.options_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.options_frame.pack(pady=10, fill="both", expand=True)

        self.buttons_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.buttons_frame.pack(pady=10, fill="x", side="bottom")

        self.next_button = ctk.CTkButton(
            self.buttons_frame,
            text="Далее",
            font=("Arial", 16),
            command=self.next_question,
            width=200,
            height=50
        )
        self.next_button.pack(side="left", padx=10)

        self.show_answer_button = ctk.CTkButton(
            self.buttons_frame,
            text="Показать ответ",
            font=("Arial", 16),
            command=self.show_correct_answer,
            width=200,
            height=50
        )
        self.show_answer_button.pack(side="right", padx=10)
        self.bind("<space>", lambda event: self.show_correct_answer())

    def show_correct_answer(self):
        correct_answer = self.questions[self.current_question]["variants"][0]
        messagebox.showinfo("Ответ", f"Правильный ответ: {correct_answer}")

    def start_timer(self):
        def update_timer():
            if self.time_remaining > 0:
                mins, secs = divmod(self.time_remaining, 60)
                timer_text = f"Оставшееся время: {mins:02d}:{secs:02d}"
                self.timer_label.configure(text=timer_text)
                self.time_remaining -= 1
                self.after(1000, update_timer)
            else:
                self.finish_test()
        update_timer()

    def show_question(self):
        question = self.questions[self.current_question]
        question_number = f"Вопрос {self.current_question + 1} из {len(self.questions)}:"
        self.question_label.configure(text=f"{question_number}\n\n{question['question']}", font=("Arial", 18, "bold"))

        # Создаём переменную выбора ответа для каждого вопроса
        self.selected_answer = ctk.IntVar(value=-1)

        # Удаляем старые кнопки 'Открыть картинку' из buttons_frame при каждом вопросе
        for widget in self.buttons_frame.winfo_children():
            if isinstance(widget, ctk.CTkButton) and widget.cget("text") == "Открыть картинку":
                widget.destroy()

        if question.get("image"):
            import os
            try:
                from PIL import Image, ImageTk
                img_path = question["image"]
                if img_path and os.path.exists(img_path):
                    img = Image.open(img_path)
                    # --- Уменьшенная копия для вопроса ---
                    # Динамически вычисляем максимальный размер картинки
                    self.update_idletasks()
                    frame_width = self.main_frame.winfo_width() or 1000
                    frame_height = self.main_frame.winfo_height() or 700
                    max_width = max(int(frame_width * 0.8), 100)
                    max_height = max(int(frame_height * 0.6), 100)
                    # Не показываем уменьшенное изображение в вопросе
                    self.image_label.pack_forget()
                    # Кнопка "Открыть картинку"
                    def open_full_image():
                        top = ctk.CTkToplevel(self)
                        top.title("Просмотр изображения")
                        top.geometry("900x700")
                        top.lift()
                        top.attributes('-topmost', True)
                        # Фрейм и Canvas с прокруткой
                        frame = ctk.CTkFrame(top)
                        frame.pack(fill="both", expand=True)
                        canvas = tk.Canvas(frame, bg="white")
                        hbar = tk.Scrollbar(frame, orient="horizontal", command=canvas.xview)
                        vbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
                        canvas.configure(xscrollcommand=hbar.set, yscrollcommand=vbar.set)
                        canvas.grid(row=0, column=0, sticky="nsew")
                        vbar.grid(row=0, column=1, sticky="ns")
                        hbar.grid(row=1, column=0, sticky="ew")
                        frame.grid_rowconfigure(0, weight=1)
                        frame.grid_columnconfigure(0, weight=1)
                        full_photo = ImageTk.PhotoImage(img)
                        img_id = canvas.create_image(0, 0, anchor="nw", image=full_photo)
                        canvas.config(scrollregion=(0, 0, img.width, img.height))
                        # Сохраняем ссылку, чтобы не удалялось
                        canvas._img_ref = full_photo
                    open_btn = ctk.CTkButton(
                        self.buttons_frame,
                        text="Открыть картинку",
                        font=("Arial", 16),
                        width=200,
                        height=50,
                        command=open_full_image
                    )
                    open_btn.pack(side="left", padx=10)
                    # Tooltip подсказка
                    def _show_tooltip(event):
                        open_btn.tooltip = tk.Toplevel(open_btn)
                        open_btn.tooltip.overrideredirect(True)
                        x = open_btn.winfo_rootx() + 60
                        y = open_btn.winfo_rooty() + 40
                        open_btn.tooltip.geometry(f'+{x}+{y}')
                        label = tk.Label(open_btn.tooltip, text="Открыть изображение в новом окне", bg="#ffffe0", relief="solid", borderwidth=1, font=("Arial", 10))
                        label.pack()
                    def _hide_tooltip(event):
                        if hasattr(open_btn, 'tooltip'):
                            open_btn.tooltip.destroy()
                    open_btn.bind("<Enter>", _show_tooltip)
                    open_btn.bind("<Leave>", _hide_tooltip)

                else:
                    self.image_label.configure(image=None, text=f"[Изображение не найдено: {img_path}]")
                    self.image_label.image = None
                    self.image_label.pack(pady=5)
            except Exception as e:
                import traceback
                error_msg = f"[Ошибка загрузки изображения: {question.get('image')}]\n{str(e)}\n{traceback.format_exc()}"
                print(error_msg)
                self.image_label.configure(image=None, text=error_msg)
                self.image_label.image = None
                self.image_label.pack(pady=5)
        else:
            self.image_label.configure(image=None, text="")
            self.image_label.image = None
            self.image_label.pack_forget()

        for widget in self.options_frame.winfo_children():
            widget.destroy()

        self.variants = question["variants"][:]
        if not self.variants:
            print(f"Ошибка: Вопрос {self.current_question + 1} не имеет вариантов ответа.")
            raise ValueError("Вопрос не имеет вариантов ответа.")

        if self.shuffle_answers:
            original_variants = self.variants[:]
            if not original_variants:  
                print(f"Ошибка: Список оригинальных вариантов пустой для вопроса {question['question']}.")
                raise IndexError("Список оригинальных вариантов пустой.")
            random.shuffle(self.variants)
            self.correct_index = self.variants.index(original_variants[0])
        else:
            self.correct_index = 0

        self.selected_answer = ctk.IntVar(value=-1)

        for i, variant in enumerate(self.variants):
            row_frame = ctk.CTkFrame(self.options_frame)
            row_frame.pack(fill="x", padx=10, pady=10)

            rb = tk.Radiobutton(
                row_frame,
                variable=self.selected_answer,
                value=i,
                font=("Arial", 50)
            )
            rb.pack(side="left", padx=10, pady=10)

            textbox = ctk.CTkTextbox(
                row_frame,
                wrap="word",
                height=100,
                width=800
            )
            textbox.insert("1.0", variant)
            textbox.configure(state="disabled", font=("Arial", 22))
            textbox.pack(side="left", fill="x", expand=True, padx=10)

    def next_question(self):
        selected_index = self.selected_answer.get()
        if selected_index == -1:
            messagebox.showwarning("Внимание", "Выберите ответ!")
            return

        is_correct = (selected_index == self.correct_index)
        self.results.append({
    "question": self.questions[self.current_question]["question"],
    "variants": self.questions[self.current_question]["variants"],
    "selected": self.variants[selected_index],
    "correct": self.variants[self.correct_index],
    "is_correct": is_correct,
    "image": self.questions[self.current_question].get("image")
})

        if is_correct:
            self.score += 4

        self.current_question += 1
        if self.current_question < len(self.questions):
            self.show_question()
        else:
            self.finish_test()

    def finish_test(self):
        self.withdraw()

        incorrect_results = [r for r in self.results if not r["is_correct"]]
        incorrect_questions = [
    {
        "question": r["question"],
        "variants": r["variants"],
        "correct_index": r["variants"].index(r["correct"]),
        "image": r.get("image")
    }
    for r in incorrect_results
]

        results_window = ctk.CTkToplevel(root)
        results_window.title("Результаты теста")
        results_window.geometry("1600x900")
        center_window(results_window, 1600, 900)

        def close_results_window():
            if messagebox.askyesno("Подтверждение", "Вы действительно хотите завершить тестирование?"):
                results_window.destroy()
                root.deiconify()

        results_window.protocol("WM_DELETE_WINDOW", close_results_window)

        results_frame = ctk.CTkFrame(
            results_window,
            corner_radius=15,
            fg_color="#ffffff"
        )
        results_frame.pack(expand=True, fill="both", padx=20, pady=20)

        # Заголовок с улучшенным стилем
        title_frame = ctk.CTkFrame(results_frame, fg_color="transparent")
        title_frame.pack(fill="x", padx=20, pady=(20, 10))

        title_label = ctk.CTkLabel(
            title_frame,
            text="Результаты теста",
            font=("Arial", 30, "bold"),
            text_color="#1a1a1a"
        )
        title_label.pack(pady=10)

        # Статистика
        correct_count = sum(r['is_correct'] for r in self.results)
        total_count = len(self.results)
        percentage = (correct_count / total_count) * 100 if total_count > 0 else 0

        stats_frame = ctk.CTkFrame(results_frame, fg_color="#f0f0f0", corner_radius=10)
        stats_frame.pack(fill="x", padx=20, pady=10)

        stats_label = ctk.CTkLabel(
            stats_frame,
            text=f"Правильных ответов: {correct_count} из {total_count} ({percentage:.1f}%)",
            font=("Arial", 24),
            text_color="#333333"
        )
        stats_label.pack(pady=15)

        # Создаем скроллируемый фрейм для результатов
        scroll_frame = ctk.CTkScrollableFrame(
            results_frame,
            fg_color="#ffffff",
            corner_radius=10
        )
        scroll_frame.pack(expand=True, fill="both", padx=20, pady=10)

        # Добавляем результаты
        for i, result in enumerate(self.results, 1):
            # Фрейм для каждого вопроса
            question_frame = ctk.CTkFrame(
                scroll_frame,
                fg_color="#f8f9fa" if result['is_correct'] else "#fff3f3",
                corner_radius=10
            )
            question_frame.pack(fill="x", padx=10, pady=5)

            # Номер вопроса
            question_number = ctk.CTkLabel(
                question_frame,
                text=f"Вопрос {i}",
                font=("Arial", 24, "bold"),
                text_color="#1a1a1a"
            )
            question_number.pack(anchor="w", padx=15, pady=(10, 5))

            # Текст вопроса
            question_text = ctk.CTkLabel(
                question_frame,
                text=result['question'],
                font=("Arial", 22),
                wraplength=1400,
                justify="left",
                text_color="#333333"
            )
            question_text.pack(anchor="w", padx=15, pady=5)

            # Ваш ответ
            answer_label = ctk.CTkLabel(
                question_frame,
                text=f"Ваш ответ: {result['selected']}",
                font=("Arial", 22),
                wraplength=1400,
                justify="left",
                text_color="#0066cc"
            )
            answer_label.pack(anchor="w", padx=15, pady=5)

            # Правильный ответ (показываем только если ответ неверный)
            if not result['is_correct']:
                correct_label = ctk.CTkLabel(
                    question_frame,
                    text=f"Правильный ответ: {result['correct']}",
                    font=("Arial", 22),
                    wraplength=1400,
                    justify="left",
                    text_color="#28a745"
                )
                correct_label.pack(anchor="w", padx=15, pady=5)

            # Статус
            status_frame = ctk.CTkFrame(
                question_frame,
                fg_color="transparent"
            )
            status_frame.pack(fill="x", padx=15, pady=(5, 10))


            status_label = ctk.CTkLabel(
                status_frame,
                text="✓ Верно" if result['is_correct'] else "✗ Неверно",
                font=("Arial", 20, "bold"),
                text_color="#28a745" if result['is_correct'] else "#dc3545"
            )
            status_label.pack(side="left")

        # Кнопки управления
        # Центрирующий контейнер для кнопок
        button_container = ctk.CTkFrame(results_frame, fg_color="transparent")
        button_container.pack(fill="x", padx=20, pady=20)

        # Внутренний фрейм для кнопок, который будет центрирован
        button_frame = ctk.CTkFrame(button_container, fg_color="transparent")
        button_frame.pack(expand=True)

        retry_all_button = ctk.CTkButton(
            button_frame,
            text="Пройти заново (все вопросы)",
            font=("Arial", 14),
            command=lambda: self.retry_with_all_questions(results_window)
        )
        retry_all_button.pack(side="left", padx=10)

        if incorrect_questions:
            retry_incorrect_button = ctk.CTkButton(
                button_frame,
                text="Пройти заново (ошибки)",
                font=("Arial", 14),
                command=lambda: self.retry_with_incorrect(incorrect_questions, results_window)
            )
            retry_incorrect_button.pack(side="left", padx=10)

            export_button = ctk.CTkButton(
                button_frame,
                text="Экспортировать неправильные ответы",
                font=("Arial", 14),
                command=lambda: export_incorrect_answers(self.results)
            )
            export_button.pack(side="left", padx=10)

            export_original_button = ctk.CTkButton(
                button_frame,
                text="Экспортировать ошибки (ориг. формат)",
                font=("Arial", 14),
                command=lambda: export_incorrect_questions_as_original_format(self.results)
            )
            export_original_button.pack(side="left", padx=10)

    def retry_with_all_questions(self, results_window):
        results_window.destroy()
        self.destroy()
        TestWindow(self.questions, self.shuffle_answers, 50)

    def retry_with_incorrect(self, incorrect_questions, results_window):
        results_window.destroy()
        self.destroy()
        TestWindow(incorrect_questions, self.shuffle_answers, 50)

    def close_results_window(self, results_window):
        if messagebox.askyesno("Подтверждение", "Вы действительно хотите завершить тестирование?"):
            results_window.destroy()
            root.deiconify()
            self.destroy()

    def close_test(self):
        if messagebox.askyesno("Подтверждение", "Вы действительно хотите завершить тестирование?"):
            root.deiconify()
            self.destroy()


main_frame = ctk.CTkFrame(root, corner_radius=15)
main_frame.pack(expand=True, fill="both", padx=20, pady=20)

header_label = ctk.CTkLabel(main_frame, text="Добро пожаловать в систему тестирования", font=("Arial", 18, "bold"))
header_label.pack(pady=10)

input_frame = ctk.CTkFrame(main_frame)
input_frame.pack(pady=10)

ctk.CTkLabel(input_frame, text="Количество вопросов:", font=("Arial", 14)).grid(row=0, column=0, padx=10, pady=10, sticky="e")

question_count_var = ctk.StringVar(value="25")
question_count_entry = ctk.CTkEntry(input_frame, textvariable=question_count_var, width=100)
question_count_entry.grid(row=0, column=1, padx=10, pady=10)

ctk.CTkLabel(input_frame, text="Время (минуты):", font=("Arial", 14)).grid(row=2, column=0, padx=10, pady=10, sticky="e")

time_limit_var = ctk.StringVar(value="50")
time_limit_entry = ctk.CTkEntry(input_frame, textvariable=time_limit_var, width=100)
time_limit_entry.grid(row=2, column=1, padx=10, pady=10)

ctk.CTkLabel(input_frame, text="Перемешивать варианты ответов?", font=("Arial", 14)).grid(row=1, column=0, padx=10, pady=10, sticky="e")

shuffle_answers_var = ctk.BooleanVar(value=True)
shuffle_checkbox = ctk.CTkCheckBox(input_frame, text="Да", variable=shuffle_answers_var)
shuffle_checkbox.grid(row=1, column=1, padx=10, pady=10)

start_button = ctk.CTkButton(main_frame, text="Начать тест", command=start_test)
start_button.pack(pady=20)


start_from_file_button = ctk.CTkButton(
    main_frame,
    text="Начать тест из файла с ошибками",
    command=start_test_with_incorrect_questions
)
start_from_file_button.pack(pady=10)


root.mainloop()